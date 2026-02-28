# academic-advisor/academic-advisor-backend/app/services/fille_parser.py
"""
File Parser Service
===================
Parses uploaded PDFs, images, text files, and documents
to extract text content and technical skills for project analysis.
"""

import os
import re
import logging
from typing import Dict, Any, List

logger = logging.getLogger(__name__)


class FileParser:
    """Parse uploaded project files to extract text and skills."""

    def __init__(self):
        self.has_pdfplumber = False
        self.has_pypdf2 = False
        self.has_pytesseract = False
        self.has_pillow = False
        self.has_docx = False
        self._check_deps()

    # ── dependency probing ──────────────────────────────────────
    def _check_deps(self):
        try:
            import pdfplumber  # noqa
            self.has_pdfplumber = True
        except ImportError:
            pass
        try:
            import PyPDF2  # noqa
            self.has_pypdf2 = True
        except ImportError:
            pass
        try:
            from PIL import Image  # noqa
            self.has_pillow = True
        except ImportError:
            pass
        try:
            import pytesseract  # noqa
            self.has_pytesseract = True
        except ImportError:
            pass
        try:
            import docx  # noqa
            self.has_docx = True
        except ImportError:
            pass

        caps = []
        if self.has_pdfplumber or self.has_pypdf2:
            caps.append("PDF")
        if self.has_pytesseract and self.has_pillow:
            caps.append("Image-OCR")
        elif self.has_pillow:
            caps.append("Image-meta")
        if self.has_docx:
            caps.append("DOCX")
        caps.append("Text/Code")
        logger.info(f"FileParser capabilities: {', '.join(caps)}")

    # ── public API ──────────────────────────────────────────────
    def parse_file(self, file_path: str, content_type: str = None) -> Dict[str, Any]:
        """Parse a single file → extracted text + skills."""
        ext = os.path.splitext(file_path)[1].lower()
        result: Dict[str, Any] = {
            "filename": os.path.basename(file_path),
            "extension": ext,
            "content_type": content_type,
            "extracted_text": "",
            "skills_found": [],
            "metadata": {},
            "parse_success": False,
            "parse_method": "none",
        }
        try:
            if ext == ".pdf":
                self._parse_pdf(file_path, result)
            elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"):
                self._parse_image(file_path, result)
            elif ext in (".doc", ".docx"):
                self._parse_docx(file_path, result)
            elif ext in (
                ".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx",
                ".java", ".cpp", ".c", ".h", ".cs", ".go", ".rs",
                ".html", ".css", ".json", ".yml", ".yaml", ".xml",
                ".sh", ".bat", ".sql", ".r", ".rb", ".kt", ".swift",
                ".ipynb",
            ):
                self._parse_text(file_path, result)
            else:
                result["metadata"]["note"] = f"Unsupported extension: {ext}"

            if result["extracted_text"]:
                result["skills_found"] = self._extract_skills(result["extracted_text"])
                result["parse_success"] = True
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {e}")
            result["metadata"]["error"] = str(e)
        return result

    def parse_multiple(self, files: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Parse many files, aggregate text + skills."""
        all_text_parts: List[str] = []
        all_skills: set = set()
        file_results: List[Dict[str, Any]] = []

        for f in files:
            r = self.parse_file(f["path"], f.get("type"))
            file_results.append(r)
            if r["extracted_text"]:
                all_text_parts.append(r["extracted_text"])
            all_skills.update(r.get("skills_found", []))

        return {
            "file_results": file_results,
            "aggregated_text": "\n".join(all_text_parts),
            "aggregated_skills": sorted(all_skills),
            "total_files": len(files),
            "successfully_parsed": sum(1 for r in file_results if r["parse_success"]),
        }

    # ── private parsers ─────────────────────────────────────────
    def _parse_pdf(self, path: str, result: Dict):
        text = ""
        if self.has_pdfplumber:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                result["metadata"]["pages"] = len(pdf.pages)
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            result["parse_method"] = "pdfplumber"
        elif self.has_pypdf2:
            import PyPDF2
            with open(path, "rb") as fh:
                reader = PyPDF2.PdfReader(fh)
                result["metadata"]["pages"] = len(reader.pages)
                for page in reader.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            result["parse_method"] = "PyPDF2"
        else:
            result["metadata"]["note"] = "No PDF library installed (pip install pdfplumber)"
            return

        result["extracted_text"] = text.strip()

    def _parse_image(self, path: str, result: Dict):
        if self.has_pillow:
            from PIL import Image
            img = Image.open(path)
            result["metadata"].update({
                "width": img.width,
                "height": img.height,
                "format": img.format,
            })

            if self.has_pytesseract:
                import pytesseract
                try:
                    text = pytesseract.image_to_string(img)
                    result["extracted_text"] = text.strip()
                    result["parse_method"] = "pytesseract_ocr"
                except Exception as e:
                    logger.warning(f"OCR failed: {e}")
                    result["parse_method"] = "image_metadata_only"
                    result["metadata"]["ocr_error"] = str(e)
            else:
                result["parse_method"] = "image_metadata_only"
                result["metadata"]["note"] = "Install pytesseract for OCR"
        else:
            result["metadata"]["note"] = "Install Pillow for image parsing"

    def _parse_text(self, path: str, result: Dict):
        try:
            with open(path, "r", errors="ignore") as fh:
                text = fh.read(100_000)  # cap at 100 KB text
            result["extracted_text"] = text
            result["parse_method"] = "text_read"
            result["metadata"]["char_count"] = len(text)
        except Exception as e:
            result["metadata"]["error"] = str(e)

    def _parse_docx(self, path: str, result: Dict):
        if not self.has_docx:
            result["metadata"]["note"] = "Install python-docx for DOCX parsing"
            return
        import docx
        doc = docx.Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        result["extracted_text"] = text
        result["parse_method"] = "python-docx"
        result["metadata"]["paragraphs"] = len(doc.paragraphs)

    # ── skill extraction ────────────────────────────────────────
    _SKILL_KEYWORDS: Dict[str, str] = {
        # Languages
        "python": "Python", "javascript": "JavaScript", "typescript": "TypeScript",
        "java ": "Java", "c++": "C++", "c#": "C#", "golang": "Go", "rust": "Rust",
        "kotlin": "Kotlin", "swift": "Swift", "dart": "Dart", "php": "PHP",
        "ruby": "Ruby", "scala": "Scala", "r programming": "R",
        # Frameworks
        "react": "React", "angular": "Angular", "vue": "Vue.js", "next.js": "Next.js",
        "django": "Django", "flask": "Flask", "fastapi": "FastAPI",
        "spring boot": "Spring Boot", "express": "Express.js", "node.js": "Node.js",
        "flutter": "Flutter", "react native": "React Native",
        # AI / ML
        "tensorflow": "TensorFlow", "pytorch": "PyTorch", "scikit-learn": "Scikit-learn",
        "keras": "Keras", "pandas": "Pandas", "numpy": "NumPy", "opencv": "OpenCV",
        "hugging face": "HuggingFace", "langchain": "LangChain",
        # Cloud / DevOps
        "docker": "Docker", "kubernetes": "Kubernetes", "aws": "AWS",
        "azure": "Azure", "gcp": "GCP", "terraform": "Terraform",
        "jenkins": "Jenkins", "github actions": "GitHub Actions", "ci/cd": "CI/CD",
        # Databases
        "mongodb": "MongoDB", "postgresql": "PostgreSQL", "mysql": "MySQL",
        "redis": "Redis", "elasticsearch": "Elasticsearch", "firebase": "Firebase",
        # Concepts
        "machine learning": "Machine Learning", "deep learning": "Deep Learning",
        "natural language processing": "NLP", "computer vision": "Computer Vision",
        "data science": "Data Science", "data analysis": "Data Analysis",
        "rest api": "REST API", "graphql": "GraphQL",
        "microservices": "Microservices", "serverless": "Serverless",
        "blockchain": "Blockchain", "iot": "IoT",
        "arduino": "Arduino", "raspberry pi": "Raspberry Pi",
        "mqtt": "MQTT", "embedded": "Embedded Systems",
    }

    def _extract_skills(self, text: str) -> List[str]:
        low = text.lower()
        found: set = set()
        for kw, canonical in self._SKILL_KEYWORDS.items():
            if kw in low:
                found.add(canonical)
        return sorted(found)


# Singleton
file_parser = FileParser()